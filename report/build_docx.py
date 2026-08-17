"""Builds report/report.docx: an editable Word version of report.html's content.

Single-column layout (Word-safe, no risky column/section XML), same figures,
tables, prompts, JSON snippets and Q&A text as the PDF. Since single-column
reflows to more than 4 pages, report.pdf (2-column, verified at exactly 4
pages) remains the file to submit; this docx exists purely for easy editing
(e.g. filling in the student name/ID) before re-exporting to PDF if desired.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "outputs" / "figures"

NAVY = RGBColor(0x0B, 0x3D, 0x5C)
GRAY = RGBColor(0x44, 0x44, 0x44)
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

doc = Document()

# ---- page setup -----------------------------------------------------------
section = doc.sections[0]
section.page_height, section.page_width = Cm(29.7), Cm(21.0)  # A4
section.top_margin = section.bottom_margin = Cm(1.8)
section.left_margin = section.right_margin = Cm(2.0)

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18


def shade_paragraph(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def border_left(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def bottom_rule(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    bottom_rule(p, "0B3D5C")
    return p


def add_subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    return p


def add_body(segments, justify=True, note=False):
    """segments: list of (text, style) where style in {'', 'b', 'i', 'c', 'q'}."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    for text, style in segments:
        r = p.add_run(text)
        r.font.size = Pt(9.5 if note else 11)
        if note:
            r.italic = True
            r.font.color.rgb = GRAY
        if "b" in style:
            r.bold = True
        if "i" in style:
            r.italic = True
        if "c" in style:
            r.font.name = MONO_FONT
            r.font.size = Pt(9.5)
        if "q" in style:
            r.bold = True
            r.font.color.rgb = NAVY
    return p


def add_code_block(text, kind="prompt"):
    color = "F4F6F8" if kind == "prompt" else "FBF7EE"
    border = "0B3D5C" if kind == "prompt" else "A06A1A"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, color)
    border_left(p, border)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line if line else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(8.5)
    return p


def add_figure(filename, caption, width_in=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    return p


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tblPr.append(borders)
    for cell in table.rows[0].cells:
        shade = OxmlElement("w:shd")
        shade.set(qn("w:val"), "clear")
        shade.set(qn("w:fill"), "EEF2F5")
        cell._tc.get_or_add_tcPr().append(shade)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def add_table(rows, col_align=None):
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if col_align and col_align[j] == "r":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# =============================================================================
# Title
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("A Hybrid Biomedical Image-Analysis Pipeline for Fluorescence-Microscopy Nuclei")
r.bold = True
r.font.size = Pt(20)
r.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(16)
r = subtitle.add_run("Assignment 3 — AI in Biomedical Imaging  |  [Student Name] — [Student ID]  |  17 August 2026")
r.font.size = Pt(11)
r.font.color.rgb = GRAY

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(12)
r = note.add_run(
    "Note: the assignment requires a PDF submission of at most 4 pages. report.pdf "
    "(2-column, verified at exactly 4 pages) is the file to submit. This .docx has the "
    "same content in a single-column layout for easy editing — update the name/ID "
    "above, then either re-export this file to PDF and check the page count, or make the "
    "same edit directly in report.html and re-render (see README)."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GRAY
shade_paragraph(note, "FFF7E6")
border_left(note, "A06A1A")

# =============================================================================
# 1. Overview
# =============================================================================
add_heading("1  Overview of methods")

add_body([
    ("Modality & data.  ", "b"),
    ("The assigned modality is synthetic fluorescence microscopy (DAPI-like nuclear staining), "
     "112 RGB 256×256 images with paired binary/instance masks, split 80/20/12 (train/val/test), "
     "plus 4 corrupted test variants for the robustness extension. All images were converted to "
     "8-bit grayscale (PIL) and confirmed at 256×256 (bilinear resize kept in the code path for "
     "portability to other sizes).", ""),
])

add_body([
    ("Pipeline, as implemented.  ", "b"),
    ("Task 1", "i"), (" pairs an EDA (sample grid, pooled intensity histogram) with a direct image "
     "description from a local vision-language model (", ""), ("llama3.2-vision", "c"),
    (" via Ollama), comparing a naive prompt against a schema-constrained, descriptive-not-diagnostic "
     "prompt, repeated 3× to expose run-to-run variability. ", ""),
    ("Task 2", "i"), (" segments the same image classically — Otsu threshold → morphological "
     "opening/closing & small-object removal → connected-component labelling → ", ""),
    ("regionprops_table", "c"),
    (" — turns the table into a short deterministic numeric summary, and sends ", ""),
    ("only that text", "i"), (" (never the pixels) to a smaller local text LLM (", ""),
    ("llama3.2", "c"), (") for a paragraph + JSON. ", ""),
    ("Task 3", "i"), (" trains a compact U-Net (4 encoder/decoder stages, base width 16, "
     "≈1.94M parameters, skip connections, combined BCE+Dice loss) for 40 epochs (Adam, "
     "lr=1e-3, batch 4, CPU) and evaluates Dice/IoU on the held-out val split. ", ""),
    ("Task 4", "i"), (" chains the trained U-Net → regionprops → code-computed JSON fields "
     "→ LLM narrative across all 12 unseen test images, aggregated to a CSV. Throughout, numeric "
     "JSON fields a user would act on are computed in code, never by LLM arithmetic — the design "
     "rationale behind this is discussed under Q4.", ""),
])

# =============================================================================
# 2. Task 1
# =============================================================================
add_heading("2  Task 1 — data preparation & VLM description")

add_figure("task1_sample_grid.png", "Fig. 1. Nine random preprocessed (grayscale, 256×256) training images.", width_in=4.6)
add_figure("task1_intensity_histogram.png",
           "Fig. 2. Pixel-intensity histogram pooled over all 80 training images: strongly right-skewed, "
           "dark background dominates, with a long low-count tail from nucleus signal — typical of "
           "fluorescence images on a black field.", width_in=5.2)

add_body([
    ("Naive prompt", "b"),
    (" (“What is this image showing? ... is there anything wrong with the tissue?”) returned "
     "free prose that volunteered an unprompted, confident clinical-sounding judgement: ", ""),
    ("“cells appear to be healthy and normal, with no visible signs of damage or abnormalities.”", "i"),
    (" Nothing in the prompt asked for a health assessment — the model defaulted to one, which is "
     "exactly the failure mode Task 1 is designed to catch.", ""),
])

add_subheading("Optimised prompt (used for all reported VLM outputs)")
add_code_block(
    "You are an assistant that provides objective, descriptive observations about a biomedical "
    "research image, for educational purposes only. You are NOT a diagnostic tool: you must not "
    "name a disease, give a diagnosis, or suggest treatment. Describe only what is visually observable.\n\n"
    "Report: modality, tissue_type, notable_features (1-2 sentences, purely descriptive), image_quality "
    "(focus/contrast/noise). If not confident about any field, write \"uncertain\" instead of guessing.\n\n"
    "Respond with ONLY a single JSON object:\n"
    '{"modality": "...", "tissue_type": "...", "notable_features": "...", "image_quality": "..."}',
    kind="prompt",
)
add_body([
    ("(Full prompt text, incl. exact schema line, in src/vlm_description.py / "
     "outputs/json_records/task1_vlm_description.json; shown condensed here for space.)", ""),
], note=True)

add_code_block(
    '{"modality": "fluorescence microscopy",\n'
    ' "tissue_type": "uncertain",\n'
    ' "notable_features": "Bright blue spots are scattered across the image,\n'
    '   possibly representing cellular structures or other biological features.",\n'
    ' "image_quality": "uncertain"}',
    kind="json",
)

add_body([
    ("Run-to-run variability.  ", "b"),
    ("Three repeats of the identical optimised prompt on the same image gave different answers for ", ""),
    ("tissue_type", "c"), (" (“uncertain”, “cell culture”, “uncertain”) and ", ""),
    ("image_quality", "c"),
    (" (“uncertain” twice, but once a full descriptive sentence). Wording of ", ""),
    ("notable_features", "c"),
    (" also changed each time. This is expected LLM sampling variance, not a bug — it is exactly "
     "why a single free-text VLM call should never be the sole record of an image (see Q1/Q4).", ""),
])

# =============================================================================
# 3. Task 2
# =============================================================================
add_heading("3  Task 2 — classical features & numbers-first interpretation")

add_body([
    ("On ", ""), ("train_004", "c"), (", Otsu + morphology detected 37 objects. A sample of the region table:", ""),
])
add_table(
    [["label", "area", "ecc.", "solidity", "mean int."],
     ["1", "267", "0.761", "0.950", "66.8"],
     ["3", "478", "0.821", "0.746", "64.8"],
     ["4", "2036", "0.826", "0.565", "73.5"],
     ["7", "102", "0.650", "0.953", "52.0"]],
    col_align=["l", "r", "r", "r", "r"],
)

add_subheading("Numbers-first interpretation prompt")
add_code_block(
    "You are an assistant that interprets a table of image-analysis measurements... You are NOT given "
    "the image itself, only these numeric measurements, and must not invent visual details... You are "
    "NOT a diagnostic tool.\n\n"
    "Measurements summary: {n_objects, area/eccentricity/solidity/mean_intensity/extent stats}\n\n"
    "1. Write ONE short paragraph (3-4 sentences) on count, density, size spread, shape regularity.\n"
    "2. Output JSON: {\"n_objects\": int, \"density_class\": \"sparse|normal|dense\",\n"
    '   "shape_regularity": "regular|irregular", "quality_flag": "ok|uncertain"}\n'
    "If ambiguous, use \"uncertain\" rather than guessing.",
    kind="prompt",
)
add_code_block(
    'LLM json:  {"n_objects": 37, "density_class": "sparse",\n'
    '            "shape_regularity": "irregular", "quality_flag": "uncertain"}\n'
    'Code json: {"n_objects": 37, "density_class": "normal",\n'
    '            "shape_regularity": "irregular", "quality_flag": "ok",\n'
    '            "mean_area": 360.9}',
    kind="json",
)
add_body([
    ("Note the disagreement: the code rule (<15 sparse, 15–40 normal, >40 dense) puts 37 objects "
     "at ", ""), ("normal", "i"),
    (" density; the LLM, given that same count in the prompt text, wrote ", ""), ("sparse", "i"),
    (" — a numeric-reasoning slip even in a numbers-only, image-free setting. This is direct "
     "evidence that the code-derived JSON, not the LLM's JSON, must be treated as ground truth "
     "(Task 4 architecture; Q4).", ""),
])

# =============================================================================
# 4. Task 3
# =============================================================================
add_heading("4  Task 3 — U-Net segmentation")
add_body([
    ("The U-Net (Section 1 architecture) was trained for 40 epochs on the 80-image train split and "
     "evaluated on the 20-image val split; Fig. 3 shows the resulting training and validation curves.", ""),
])
add_figure("task3_loss_dice_curves.png",
           "Fig. 3. Training loss (BCE+Dice) and validation Dice/IoU over 40 epochs. Both curves plateau "
           "by ~epoch 15-20; the visible dip in early-epoch Dice/IoU (epochs 1-3) reflects the "
           "randomly-initialised decoder before it learns to close small holes.", width_in=6.2)

add_table(
    [["Split", "Method", "Mean Dice", "Mean IoU"],
     ["val (n=20)", "Otsu (classical)", "0.976", "0.953"],
     ["val (n=20)", "U-Net (BCE+Dice, 40 ep.)", "0.997", "0.994"]],
    col_align=["l", "l", "r", "r"],
)

add_figure("task3_prediction_panels.png",
           "Fig. 4. Input / ground-truth / U-Net prediction for 3 validation images (val_000–val_002) "
           "— predicted masks are visually close to pixel-identical to ground truth.", width_in=5.4)

# =============================================================================
# 5. Otsu vs UNet + Task 4
# =============================================================================
add_heading("5  U-Net vs Otsu, and the hybrid pipeline (Task 4)")

add_figure("otsu_vs_unet_panels.png",
           "Fig. 5. Qualitative comparison on the two extreme cases from val: U-Net's largest margin "
           "(val_019, touching/medium-contrast nuclei — Otsu leaves ragged, notched edges) and its "
           "smallest margin (val_002, large well-separated bright nuclei — a global threshold is "
           "already nearly correct).", width_in=6.2)

add_body([
    ("The U-Net matched or beat Otsu on ", ""), ("every", "i"),
    (" validation image; there was no case where Otsu was actually better, only cases where its "
     "disadvantage was small. That the gap is largest on medium-contrast/touching objects and smallest "
     "on large, bright, well-isolated ones (Fig. 5) is consistent with Otsu's single global threshold "
     "being a good approximation only when foreground/background separation is already easy.", ""),
])

add_subheading("Task 4 example record (test_000, unseen)")
add_code_block(
    '{"image_id": "test_000", "n_objects": 8, "mean_area": 191.0,\n'
    ' "density_class": "sparse", "quality_flag": "ok",\n'
    ' "narrative": "The test_000 image contains 8 objects with varying sizes,\n'
    '   as indicated by a standard deviation of 111 pixels in their areas.\n'
    '   These objects appear to have a range of shapes ... suggesting some\n'
    '   elongation or irregularity. However, the mean solidity and extent\n'
    '   values suggest that most ... are roughly convex and well-defined ..."}',
    kind="json",
)
add_body([
    ("All 12 test images were processed this way (U-Net mask → regionprops → code JSON fields "
     "→ LLM narrative) and aggregated into ", ""), ("outputs/task4_hybrid_records.csv", "c"),
    (" (one row per image; ", ""), ("n_objects", "c"), (" ranged 8–43, ", ""),
    ("density_class", "c"),
    (" sparse/normal/dense throughout). The narrative prompt keeps Task 2's numbers-only, "
     "non-diagnostic framing but asks for prose only (the JSON fields are already fixed by code, so "
     "the LLM is not asked to repeat them):", ""),
])
add_code_block(
    "You are an assistant that writes a short, objective narrative for a biomedical research image, "
    "for educational purposes only. You are NOT a diagnostic tool. You are given only numeric "
    "measurements below, not the image itself -- do not invent visual details beyond what the numbers "
    "imply.\n"
    "Image id: {image_id}\n"
    "Measurements from automatic segmentation: {summary_text}\n"
    "Write ONE short paragraph (3-4 sentences) ... If the numbers are ambiguous, say so rather than "
    "guessing. Plain text only, no JSON.",
    kind="prompt",
)

add_subheading("Robustness extension (bonus)")
add_figure("robustness_panels.png",
           "Fig. 6. Corruption propagation on test_000. Blur softens object boundaries but the mask "
           "still finds discrete objects. Low contrast is catastrophic: the U-Net mask becomes 100% "
           "foreground (the whole frame), silently breaking every downstream stage.", width_in=6.2)
add_body([
    ("Two corrupted variants (heavy Gaussian blur; low-contrast rescale) of 2 test images were traced "
     "through all four stages. ", ""), ("Blur", "b"),
    (": pixel std drops ~40–44% (stage 1), mask Dice-vs-GT collapses from ~0.997 to ~0.65 "
     "(stage 2, Fig. 6 middle column), object count drops modestly. ", ""), ("Low-contrast", "b"),
    (": pixel mean shifts by hundreds of percent (stage 1), the mask ", ""), ("completely fails", "i"),
    (" to 100% foreground (Dice≈0.05–0.36, Fig. 6 right column), and the region table "
     "collapses to a single “object” with ", ""), ("area=65535", "c"),
    (" — exactly 256×256. Critically, the LLM narrative for this case still reads fluently "
     "(“a single, elongated oval-shaped object…”) with no hint of failure — only the "
     "structured field ", ""), ("area=65535", "c"), (" exposes it. The earliest stage at which the "
     "corruption is ", ""), ("detectable at all", "i"), (" is stage 1 (raw pixel statistics, before "
     "segmentation even runs); the earliest stage at which it becomes ", ""),
    ("unambiguous without instrumenting raw pixels", "i"),
    (" is the stage-3 feature table, not the narrative.", ""),
])

# =============================================================================
# 6. Discussion / Q&A
# =============================================================================
add_heading("6  Discussion & answers to the set questions")

add_body([
    ("Q1 — VLM vs numbers-first: which is more useful, which is more trustworthy?  ", "q"),
    ("The direct VLM description (Task 1) is more ", ""), ("useful", "i"),
    (" for a fast, human-readable first impression — it reads naturally and needs no downstream "
     "code to consume. It is less ", ""), ("trustworthy", "i"),
    (": it is non-deterministic (Section 2), it is not grounded in any measurement that can be "
     "independently re-checked, and the naive-prompt run showed it will volunteer confident-sounding "
     "judgements it was never asked for. The numbers-first pipeline (Task 2) is more trustworthy "
     "because every quantity is reproducible — re-running Otsu+regionprops on the same image "
     "always gives the same table — and the LLM's role is reduced to phrasing rather than "
     "measuring. It is not perfectly trustworthy either (Section 3 showed a density-class "
     "misclassification), which is exactly why the ", ""), ("code-derived", "i"),
    (" JSON, not either LLM's JSON, is treated as ground truth throughout.", ""),
])

add_body([
    ("Q2 — Did the U-Net improve on Otsu? Examples.  ", "q"),
    ("Yes, on every validation image (mean Dice 0.997 vs 0.976; mean IoU 0.994 vs 0.953). U-Net's "
     "advantage is largest on ", ""), ("val_019", "c"),
    (" (Otsu 0.973 vs U-Net 0.999) where nuclei are medium-contrast and partly touching — Otsu's "
     "fixed global threshold leaves ragged edges and small spurious fragments (Fig. 5, top row). Otsu "
     "is closest to U-Net (though still behind) on ", ""), ("val_002", "c"),
    (" (0.982 vs 0.997), where nuclei are large, bright, and well-isolated against a clean dark "
     "background — conditions under which a single threshold is nearly sufficient by "
     "construction.", ""),
])

add_body([
    ("Q3 — Dice/IoU meaning and failure modes.  ", "q"),
    ("Dice = 2|P∩G| / (|P|+|G|) and IoU = |P∩G| / |P∪G| both measure pixel-set overlap "
     "between predicted (P) and ground-truth (G) foreground; IoU penalises boundary/size mismatches "
     "more harshly than Dice for the same error. Mean val Dice 0.997 / IoU 0.994 mean predicted masks "
     "are almost pixel-identical to ground truth on this dataset. The (narrow, 0.995–0.999) "
     "per-image range is lowest on images with more, smaller, or more closely packed nuclei (e.g. "
     "val_004, val_005, val_012) — consistent with the qualitative panels, where residual error "
     "concentrates at object boundaries and at touching instances. Because the model outputs one "
     "binary foreground mask rather than per-instance labels, it structurally cannot separate two "
     "touching nuclei into two objects even when the pixel-level Dice looks excellent — a real "
     "limitation Dice/IoU do not expose.", ""),
])

add_body([
    ("Q4 — Where can the LLM hallucinate, and what mitigates it?  ", "q"),
    ("Three concrete points were observed: (i) Task 1's free-text ", ""), ("tissue_type", "c"),
    ("/", ""), ("image_quality", "c"),
    (" fields changing answer across identical re-runs; (ii) Task 2's LLM assigning "
     "density_class=\"sparse\" to 37 objects when the code rule (and its own stated count) implies "
     "“normal”; (iii) the robustness test's narrative describing a physically broken, "
     "whole-frame mask as “a single, elongated oval-shaped object” without flagging anything "
     "wrong. Mitigations used in the code: a constrained JSON schema that explicitly permits "
     "“uncertain” (removing the pressure to fabricate a confident answer); an explicit "
     "“not a diagnostic tool” role instruction; isolating the interpretation LLM from pixels "
     "entirely in Tasks 2/4 so it cannot invent visual detail beyond the numbers it is given; a "
     "regex-based JSON-extraction-with-retry helper (", ""), ("llm_utils.chat_json", "c"),
    (") that rejects and retries malformed output; and, most importantly, computing every numeric "
     "field a downstream user would act on (", ""), ("n_objects, mean_area, density_class, quality_flag", "c"),
    (") directly from the regionprops table in code, never from the LLM's own arithmetic. Keeping "
     "this code-derived JSON as the source of truth matters because it is deterministic and diffable "
     "(identical input → identical output, unlike the LLM text), can be automatically "
     "range-checked (e.g. flag ", ""), ("area", "c"), (" ≥ image area, exactly what would have "
     "caught the corrupted-image failure), and lets an auditor trace any number back to the exact "
     "pixels and function that produced it — the fluent narrative alone cannot offer any of "
     "that.", ""),
])

add_body([
    ("Q5 — Clinical trust and the highest-impact change.  ", "q"),
    ("No part of this system should be trusted in a real clinical setting yet. The near-perfect Dice "
     "(0.997) reflects a fully synthetic dataset of clean, well-separated, uniformly-stained blobs on "
     "a flat dark background — the model has never seen real staining variability, overlapping "
     "nuclei at density, or acquisition artefacts, so the number says more about the dataset's ease "
     "than about the model's real-world skill. The VLM step is demonstrably non-deterministic and can "
     "produce confident-sounding but unverified claims. And evaluation is at the pixel/mask level only "
     "— there is no instance-level check, so the pipeline could report a plausible object count "
     "for the wrong reason. Considering accuracy, auditability, and dataset limits together, the "
     "single highest-impact change would be replacing the synthetic dataset with a real, "
     "expert-annotated microscopy/histology set (with a held-out clinical-site split) ", ""),
    ("and", "i"),
    (" wiring the automatic range/plausibility checks the robustness test motivates (e.g. reject or "
     "flag any code-derived record whose ", ""), ("area", "c"),
    (" is implausible, before any narrative reaches a user) — together these close both the "
     "generalisation gap and the “plausible but wrong” narrative failure mode this report "
     "demonstrates.", ""),
])

# =============================================================================
# References
# =============================================================================
add_heading("References")
refs = [
    "Ronneberger, O., Fischer, P., Brox, T. (2015). U-Net: Convolutional Networks for Biomedical Image Segmentation. MICCAI.",
    "Otsu, N. (1979). A Threshold Selection Method from Gray-Level Histograms. IEEE Trans. Systems, Man, and Cybernetics, 9(1), 62–66.",
    "van der Walt, S. et al. (2014). scikit-image: image processing in Python. PeerJ, 2:e453.",
    "Sørensen, T. (1948). A method of establishing groups of equal amplitude... Biol. Skr., 5, 1–34. (Dice/Sørensen coefficient.)",
    "Meta AI (2024). Llama 3.2 Model Card (incl. vision variants). Served locally via Ollama (ollama.com).",
    "Caicedo, J. C. et al. (2019). Nucleus segmentation across imaging experiments: the 2018 Data Science Bowl. Nature Methods, 16, 1247–1253.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    r.font.size = Pt(9.5)

out_path = Path(__file__).resolve().parent / "report.docx"
doc.save(str(out_path))
print("saved", out_path)
